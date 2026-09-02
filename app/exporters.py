from __future__ import annotations

import html
import base64
import json
import os
import subprocess
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path

from app.paths import RESOURCE_ROOT, USER_FONT_DIR
from app.media_prompts import media_prompts_markdown


ROOT = RESOURCE_ROOT
FONT_DIR = RESOURCE_ROOT / "fonts"


def sermon_with_media_prompts(sermon: str, meta: dict) -> str:
    appendix = media_prompts_markdown(meta.get("media_prompts") if isinstance(meta, dict) else None)
    return f"{sermon.rstrip()}\n\n{appendix}" if appendix else sermon


def _font_path(filename: str) -> Path:
    user_path = USER_FONT_DIR / filename
    return user_path if user_path.exists() else FONT_DIR / filename


def _embedded_font_face(family: str, filename: str, weight: int, mime: str = "font/woff2") -> str:
    path = _font_path(filename)
    source = f"local('{family}')"
    if not path.exists() and path.suffix.lower() == ".woff2":
        ttf = _font_path(Path(filename).with_suffix(".ttf").name)
        if ttf.exists():
            path = ttf; mime = "font/ttf"
    if path.exists():
        encoded = base64.b64encode(path.read_bytes()).decode("ascii")
        source += f",url(data:{mime};base64,{encoded})"
    return f"@font-face{{font-family:'{family}';src:{source};font-weight:{weight};font-style:normal;font-display:swap}}"


def dashboard_font_css() -> str:
    return "".join([
        _embedded_font_face("S-Core Dream", "S-CoreDream-3Light.woff2", 300),
        _embedded_font_face("S-Core Dream", "S-CoreDream-5Medium.woff2", 500),
        _embedded_font_face("S-Core Dream", "S-CoreDream-6Bold.woff2", 700),
    ])


def dashboard_html(*, sermon: str, meta: dict, sources: list[dict]) -> str:
    sermon = sermon_with_media_prompts(sermon, meta)
    esc = lambda value: html.escape(str(value or ""))
    cards = "".join(
        f'''<article class="source"><div><b>{esc(s['translation'])}</b><span>{esc(s['language'])}</span></div>
        <h3>{esc(s['reference'])}</h3><p>{esc(s['text'])}</p><small>{esc(s.get('license_note') or '사용조건 기록 없음')}</small></article>'''
        for s in sources
    ) or '<p>검색된 성경 근거가 없습니다.</p>'
    unchecked = meta.get("unchecked_references", [])
    check_text = "참조 1차 확인 완료" if not unchecked else f"DB 미확인 참조 {len(unchecked)}건"
    check_class = "good" if not unchecked else "warn"
    generated = datetime.now().strftime("%Y-%m-%d %H:%M")
    notes = meta.get("original_notes", []) if isinstance(meta.get("original_notes", []), list) else []
    audit = meta.get("audit", {}) if isinstance(meta.get("audit", {}), dict) else {}
    review = meta.get("review_state", {}) if isinstance(meta.get("review_state", {}), dict) else {}
    audit_status = audit.get("status") or "기록 없음"
    review_status = review.get("state") or "미검토"
    audit_warnings = audit.get("warnings", []) if isinstance(audit.get("warnings", []), list) else []
    audit_detail = " · ".join(str(w) for w in audit_warnings) or "자동 점검 경고 없음"
    citation = meta.get("citation_analysis") if isinstance(meta.get("citation_analysis"), dict) else audit.get("citation_analysis", {})
    citation = citation if isinstance(citation, dict) else {}
    citation_cards = "".join(
        f'''<article class="source"><div><b>문장 {esc(x.get('sentence'))} · 근거 연결</b><span>{esc(', '.join(x.get('references', [])))}</span></div><p>{esc(x.get('text'))}</p></article>'''
        for x in citation.get("mappings", [])[:30]
    )
    citation_cards += "".join(
        f'''<article class="source warn"><div><b>문장 {esc(x.get('sentence'))} · 확인 필요</b><span>{esc(x.get('reason'))}</span></div><p>{esc(x.get('text'))}</p></article>'''
        for x in citation.get("unsupported_claims", [])[:30]
    )
    citation_cards = citation_cards or '<p>문장별 근거 분석 기록이 없습니다.</p>'
    footer_text = "목회자 승인 · 최종 잠금본" if review_status == "locked" else "목회자 검토 전/진행 중 초안"
    revision_note = f"AI 수정 제안 반영본 · 상위 버전 v{meta.get('revision_parent_version')} · 제안 {len(meta.get('applied_suggestion_ids', []))}건" if meta.get("revision_parent_version") else "AI 수정 제안 반영 이력 없음"
    project = meta.get("project", {}) if isinstance(meta.get("project"), dict) else {}
    note_cards = "".join(
        f'''<article class="source"><div><b>{esc(n.get('language'))} · {esc(n.get('lemma'))}</b><span>{esc(n.get('transliteration'))}</span></div>
        <p><b>뜻</b> {esc(n.get('gloss'))}<br><b>형태</b> {esc(n.get('morphology'))}</p><small>{esc(n.get('source'))} · {esc(n.get('license_note'))}</small></article>'''
        for n in notes
    ) or '<p>등록된 원어 어휘 근거가 없습니다.</p>'
    return f'''<!doctype html><html lang="ko"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{esc(meta.get('topic') or '설교문')} · 설교 대시보드</title><style>
{dashboard_font_css()}:root{{--ink:#17231f;--green:#245d50;--paper:#f4f7f5;--line:#dbe4df;--warm:#fff4d8}}*{{box-sizing:border-box}}body{{margin:0;font-family:'S-Core Dream','Noto Sans KR','Malgun Gothic','Apple SD Gothic Neo',sans-serif;color:var(--ink);background:var(--paper);line-height:1.72}}
.hero{{background:linear-gradient(135deg,#173f37,#327262);color:white;padding:52px max(24px,calc((100vw - 1120px)/2))}}.kicker{{letter-spacing:.16em;font-size:12px;color:#bde1d5}}h1{{font-size:clamp(30px,5vw,48px);line-height:1.25;margin:7px 0}}.hero p{{color:#dceee8}}
.wrap{{max-width:1120px;margin:26px auto;padding:0 20px}}.stats{{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-top:-48px}}.stat{{background:white;border:1px solid var(--line);border-radius:16px;padding:18px;box-shadow:0 9px 30px #173f3712}}.stat small{{display:block;color:#6a7b75}}.stat b{{font-size:21px;color:var(--green)}}.good b{{color:#28734e}}.warn{{background:var(--warm)}}
.panel{{background:white;border:1px solid var(--line);border-radius:18px;padding:28px;margin:18px 0}}.panel h2{{margin:0 0 16px;font-size:20px}}.sermon{{white-space:pre-wrap;font-size:16px}}.sources{{display:grid;grid-template-columns:repeat(2,1fr);gap:12px}}.source{{background:#f8fbf9;border:1px solid var(--line);border-radius:13px;padding:16px}}.source div{{display:flex;justify-content:space-between}}.source span,.source small{{color:#6b7a75;font-size:12px}}.source p{{white-space:pre-wrap}}.footer{{text-align:center;color:#687872;font-size:12px;padding:22px}}
@media(max-width:760px){{.stats,.sources{{grid-template-columns:1fr 1fr}}}}@media(max-width:520px){{.stats,.sources{{grid-template-columns:1fr}}}}@media print{{body{{background:white}}.hero{{padding:24px;color:#17231f;background:white;border-bottom:2px solid #245d50}}.hero p,.kicker{{color:#555}}.stats{{margin-top:12px}}.panel,.stat{{box-shadow:none}}}}
</style></head><body><header class="hero"><div class="kicker">SERMON EVIDENCE DASHBOARD</div><h1>{esc(meta.get('topic') or '설교문')}</h1><p>{esc(meta.get('main_reference') or '중심본문 자동 검색')} · {esc(meta.get('audience'))} · {esc(meta.get('tradition'))}</p><p>{esc(project.get('service_date'))} · {esc(project.get('series_name'))} · {esc(project.get('preacher'))}</p></header>
<main class="wrap"><section class="stats"><div class="stat"><small>목표 시간</small><b>{esc(meta.get('target_minutes'))}분</b></div><div class="stat"><small>예상 낭독</small><b>{esc(meta.get('minutes_estimate'))}분</b></div><div class="stat"><small>생성 감사</small><b>{esc(audit_status)}</b><small>{esc(audit_detail)}</small></div><div class="stat {check_class}"><small>목회자 검토</small><b>{esc(review_status)}</b><small>{esc(check_text)}</small></div></section>
<section class="panel"><h2>설교 원고</h2><p><small>{esc(revision_note)}</small></p><div class="sermon">{esc(sermon)}</div></section><section class="panel"><h2>문장별 성경 근거 연결</h2><p>명시 근거 {esc(citation.get('mapped_count',0))}건 · 확인 필요 {esc(citation.get('unsupported_count',0))}건</p><div class="sources">{citation_cards}</div></section><section class="panel"><h2>사용한 성경 근거</h2><div class="sources">{cards}</div></section><section class="panel"><h2>히브리어·헬라어 원어 근거</h2><div class="sources">{note_cards}</div></section></main><footer class="footer">LM Studio 로컬 AI · {esc(footer_text)} · 생성 {generated}</footer></body></html>'''


def _pdf_font_css() -> str:
    candidates = [
        ("NanumSquare", "NanumSquareR.ttf", 400), ("NanumSquare", "NanumSquareB.ttf", 700),
        ("NanumGothic", "NanumGothic.ttf", 400), ("NanumGothic", "NanumGothicBold.ttf", 700),
    ]
    rules = []
    for family, filename, weight in candidates:
        path = _font_path(filename)
        if path.exists():
            rules.append(f"@font-face{{font-family:'{family}';src:url('{path.resolve().as_uri()}');font-weight:{weight}}}")
    return "".join(rules)


def _pdf_local_font() -> tuple[Path | None, Path | None, str]:
    """Return a usable Korean TTF/OTF pair, preferring the requested Nanum fonts."""
    for regular, bold, family in (
        ("NanumSquareR.ttf", "NanumSquareB.ttf", "NanumSquare"),
        ("NanumGothic.ttf", "NanumGothicBold.ttf", "NanumGothic"),
    ):
        path = _font_path(regular)
        if path.exists():
            bold_path = _font_path(bold)
            return path, bold_path if bold_path.exists() else path, family
    if os.name == "nt":
        roots = [Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts"]
        local = os.environ.get("LOCALAPPDATA")
        if local:
            roots.append(Path(local) / "Microsoft" / "Windows" / "Fonts")
        for root in roots:
            if not root.exists():
                continue
            nanum = next(iter(list(root.glob("*NanumGothic*.ttf")) + list(root.glob("*NanumSquare*.ttf"))), None)
            if nanum:
                return nanum, nanum, "NanumGothic"
            malgun = root / "malgun.ttf"
            if malgun.exists():
                bold = root / "malgunbd.ttf"
                return malgun, bold if bold.exists() else malgun, "Malgun Gothic"
        return None, None, ""
    try:
        result = subprocess.run(["fc-match", "-f", "%{file}", "NanumGothic"], capture_output=True, text=True, timeout=3)
        path = Path(result.stdout.strip())
        return (path, path, "NanumGothic") if path.exists() and "Nanum" in path.name else (None, None, "")
    except (OSError, subprocess.SubprocessError):
        return None, None, ""


def _nanum_available() -> bool:
    # Kept as the historical helper name; Windows Malgun Gothic is an accepted
    # local-system fallback when the preferred Nanum family is unavailable.
    return _pdf_local_font()[0] is not None


def _pdf_engine_status() -> tuple[bool, str, str]:
    errors = []
    if os.name == "nt":
        try:
            import reportlab  # noqa: F401
            return True, "reportlab", ""
        except Exception as exc:
            errors.append(f"ReportLab: {exc}")
    try:
        from weasyprint import HTML  # noqa: F401
        return True, "weasyprint", ""
    except Exception as exc:
        errors.append(f"WeasyPrint: {exc}")
    try:
        import reportlab  # noqa: F401
        return True, "reportlab", " · ".join(errors)
    except Exception as exc:
        errors.append(f"ReportLab: {exc}")
    return False, "", " · ".join(errors)


def pdf_environment_status() -> dict:
    """PDF 출력 전 필요한 엔진/한글 글꼴 상태를 파일 생성 없이 점검한다."""
    regular_font, _bold_font, font_family = _pdf_local_font()
    font_ready = regular_font is not None
    engine_ready, engine, engine_error = _pdf_engine_status()
    return {
        "ready": bool(font_ready and engine_ready),
        "engine_ready": engine_ready,
        "font_ready": font_ready,
        "engine": engine,
        "font_family": font_family,
        "engine_error": engine_error,
    }


def pdf_document_html(*, sermon: str, meta: dict, sources: list[dict]) -> str:
    sermon = sermon_with_media_prompts(sermon, meta)
    esc = lambda v: html.escape(str(v or ""))
    notes = meta.get("original_notes", []) if isinstance(meta.get("original_notes", []), list) else []
    evidence = "".join(f"<tr><td>{esc(s.get('reference'))}</td><td>{esc(s.get('translation'))}</td><td>{esc(s.get('text'))}</td></tr>" for s in sources[:24])
    note_rows = "".join(f"<tr><td>{esc(n.get('language'))}</td><td>{esc(n.get('lemma'))}</td><td>{esc(n.get('transliteration'))}</td><td>{esc(n.get('gloss'))}</td><td>{esc(n.get('source'))}</td></tr>" for n in notes)
    audit = meta.get("audit", {}) if isinstance(meta.get("audit", {}), dict) else {}
    review = meta.get("review_state", {}) if isinstance(meta.get("review_state", {}), dict) else {}
    project = meta.get("project", {}) if isinstance(meta.get("project"), dict) else {}
    citation = meta.get("citation_analysis") if isinstance(meta.get("citation_analysis"), dict) else audit.get("citation_analysis", {})
    citation = citation if isinstance(citation, dict) else {}
    citation_rows = "".join(
        f"<tr><td>{esc(x.get('sentence'))}</td><td>{esc(', '.join(x.get('references', [])))}</td><td>연결</td><td>{esc(x.get('text'))}</td></tr>"
        for x in citation.get("mappings", [])[:30]
    ) + "".join(
        f"<tr><td>{esc(x.get('sentence'))}</td><td>-</td><td>확인 필요</td><td>{esc(x.get('text'))}</td></tr>"
        for x in citation.get("unsupported_claims", [])[:30]
    )
    return f'''<!doctype html><html lang="ko"><head><meta charset="utf-8"><style>
{_pdf_font_css()}@page{{size:A4 landscape;margin:14mm 16mm 14mm 16mm;@bottom-right{{content:"Page " counter(page);font-size:8pt;color:#66736e}}}}
body{{font-family:'NanumSquare','NanumGothic','Noto Sans KR','Malgun Gothic','Apple SD Gothic Neo',sans-serif;font-size:10.5pt;line-height:1.65;color:#19241f}}h1{{font-size:23pt;color:#245d50;margin:0 0 4mm}}h2{{font-size:14pt;color:#245d50;margin:7mm 0 3mm}}.meta{{display:flex;gap:5mm;background:#eef5f1;padding:3mm 4mm;border-radius:3mm}}.sermon{{white-space:pre-wrap}}table{{width:100%;border-collapse:collapse;font-size:8.8pt}}th{{background:#245d50;color:white}}th,td{{padding:2.2mm;border:0.3mm solid #d9e2de;vertical-align:top}}.notice{{color:#697771;font-size:8.5pt}}
</style></head><body><h1>{esc(meta.get('topic') or '설교문')}</h1><div class="meta"><span>중심본문: {esc(meta.get('main_reference'))}</span><span>예배일: {esc(project.get('service_date'))}</span><span>목표: {esc(meta.get('target_minutes'))}분</span><span>예상: {esc(meta.get('minutes_estimate'))}분</span><span>속도: {esc(meta.get('reading_cpm') or 330)}자/분</span><span>감사: {esc(audit.get('status') or '없음')}</span><span>검토: {esc(review.get('state') or '미검토')}</span><span>AI 수정: {esc('v'+str(meta.get('revision_parent_version'))+' 기반' if meta.get('revision_parent_version') else '없음')}</span></div>
<h2>설교 원고</h2><div class="sermon">{esc(sermon)}</div><h2>문장별 성경 근거 연결</h2><table><thead><tr><th>문장</th><th>참조</th><th>상태</th><th>문장 내용</th></tr></thead><tbody>{citation_rows}</tbody></table><h2>성경 근거</h2><table><thead><tr><th>참조</th><th>번역/자료</th><th>본문</th></tr></thead><tbody>{evidence}</tbody></table>
<h2>히브리어·헬라어 원어 근거</h2><table><thead><tr><th>언어</th><th>원어</th><th>음역</th><th>뜻</th><th>출처</th></tr></thead><tbody>{note_rows}</tbody></table><p class="notice">{esc('목회자 승인 최종 잠금본' if review.get('state') == 'locked' else '목회자 검토 전/진행 중 AI 초안')}</p></body></html>'''


def write_pdf(path: Path, *, sermon: str, meta: dict, sources: list[dict]) -> None:
    regular_font, bold_font, _family = _pdf_local_font()
    if regular_font is None:
        raise RuntimeError("한글 PDF 글꼴을 찾지 못했습니다. NanumSquare/NanumGothic TTF를 fonts 폴더에 넣으세요. Windows에서는 설치된 맑은 고딕도 자동 사용합니다.")
    if os.name == "nt":
        _write_pdf_reportlab(path, sermon=sermon, meta=meta, sources=sources,
                             regular_font=regular_font, bold_font=bold_font or regular_font)
        return
    try:
        from weasyprint import HTML
        HTML(string=pdf_document_html(sermon=sermon, meta=meta, sources=sources), base_url=str(ROOT)).write_pdf(path)
        return
    except Exception:
        # WeasyPrint needs native Pango/GTK libraries on some Windows setups.
        # ReportLab is the native-library-independent fallback installed with
        # requirements-pdf.txt, so PDF remains usable without MSYS2/Pango.
        _write_pdf_reportlab(path, sermon=sermon, meta=meta, sources=sources,
                             regular_font=regular_font, bold_font=bold_font or regular_font)


def _write_pdf_reportlab(path: Path, *, sermon: str, meta: dict, sources: list[dict],
                         regular_font: Path, bold_font: Path) -> None:
    sermon = sermon_with_media_prompts(sermon, meta)
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except Exception as exc:
        raise RuntimeError("PDF 엔진을 사용할 수 없습니다. requirements-pdf.txt를 다시 설치하세요.") from exc

    pdfmetrics.registerFont(TTFont("SermonKorean", str(regular_font)))
    pdfmetrics.registerFont(TTFont("SermonKoreanBold", str(bold_font)))
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("SermonBody", parent=styles["BodyText"], fontName="SermonKorean",
                            fontSize=10.5, leading=17, spaceAfter=5)
    heading = ParagraphStyle("SermonHeading", parent=normal, fontName="SermonKoreanBold",
                             fontSize=14, leading=20, textColor=colors.HexColor("#245d50"), spaceBefore=10, spaceAfter=6)
    title = ParagraphStyle("SermonTitle", parent=heading, fontSize=22, leading=28, alignment=TA_CENTER, spaceAfter=10)
    small = ParagraphStyle("SermonSmall", parent=normal, fontSize=8.5, leading=13, textColor=colors.HexColor("#66736e"))
    esc = lambda value: html.escape(str(value or "")).replace("\n", "<br/>")
    story = [Paragraph(esc(meta.get("topic") or "설교문"), title)]
    story.append(Paragraph(f"중심본문: {esc(meta.get('main_reference'))} · 목표: {esc(meta.get('target_minutes'))}분", small))
    story.extend([Spacer(1, 4 * mm), Paragraph("설교 원고", heading)])
    for block in sermon.split("\n\n"):
        if block.strip():
            story.append(Paragraph(esc(block.strip()), normal))
    if sources:
        story.append(Paragraph("성경 근거", heading))
        rows = [[Paragraph("참조", small), Paragraph("번역/자료", small), Paragraph("본문", small)]]
        for source in sources[:24]:
            rows.append([Paragraph(esc(source.get("reference")), small), Paragraph(esc(source.get("translation")), small), Paragraph(esc(source.get("text")), small)])
        table = Table(rows, colWidths=[34 * mm, 34 * mm, 185 * mm], repeatRows=1)
        table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), .3, colors.HexColor("#d9e2de")),
                                   ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef5f1")),
                                   ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 5),
                                   ("RIGHTPADDING", (0, 0), (-1, -1), 5), ("TOPPADDING", (0, 0), (-1, -1), 4),
                                   ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
        story.append(table)
    path.parent.mkdir(parents=True, exist_ok=True)
    SimpleDocTemplate(str(path), pagesize=landscape(A4), leftMargin=16 * mm, rightMargin=16 * mm,
                      topMargin=14 * mm, bottomMargin=14 * mm, title=str(meta.get("topic") or "설교문")).build(story)


def write_docx(path: Path, *, sermon: str, meta: dict) -> None:
    sermon = sermon_with_media_prompts(sermon, meta)
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except Exception as exc:
        raise RuntimeError("Word 출력 모듈 python-docx가 설치되지 않았습니다.") from exc
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1); section.bottom_margin = Inches(1); section.left_margin = Inches(1); section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]; normal.font.name = "NanumGothic"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "NanumGothic"); normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8); normal.paragraph_format.line_spacing = 1.333
    for name, size, color, before, after in [("Heading 1",16,"2E74B5",18,10),("Heading 2",13,"2E74B5",12,6),("Heading 3",12,"1F4D78",8,4)]:
        st=styles[name]; st.font.name="NanumGothic"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"NanumGothic"); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)
    title = doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; title.paragraph_format.space_after=Pt(8)
    run=title.add_run(str(meta.get("topic") or "설교문")); run.font.name="NanumGothic"; run._element.rPr.rFonts.set(qn("w:eastAsia"),"NanumGothic"); run.font.size=Pt(26); run.bold=True; run.font.color.rgb=RGBColor(36,93,80)
    audit = meta.get("audit", {}) if isinstance(meta.get("audit", {}), dict) else {}
    review = meta.get("review_state", {}) if isinstance(meta.get("review_state", {}), dict) else {}
    project = meta.get("project", {}) if isinstance(meta.get("project", {}), dict) else {}
    revision_text = f" · AI 수정 v{meta.get('revision_parent_version')} 기반" if meta.get("revision_parent_version") else ""
    sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER; sub.paragraph_format.space_after=Pt(20); sub.add_run(f"{meta.get('main_reference','')} · {meta.get('tradition','')} · {project.get('service_date','')} · {project.get('preacher','')} · 약 {meta.get('target_minutes','20')}분 · {meta.get('reading_cpm',330)}자/분 · 감사 {audit.get('status','없음')} · 검토 {review.get('state','미검토')}{revision_text}")
    for raw in sermon.splitlines():
        line=raw.strip()
        if line.startswith("### "): doc.add_paragraph(line[4:], style="Heading 3")
        elif line.startswith("## "): doc.add_paragraph(line[3:], style="Heading 2")
        elif line.startswith("# "): doc.add_paragraph(line[2:], style="Heading 1")
        elif line: doc.add_paragraph(line)
    citation = meta.get("citation_analysis") if isinstance(meta.get("citation_analysis"), dict) else audit.get("citation_analysis", {})
    citation = citation if isinstance(citation, dict) else {}
    doc.add_paragraph("문장별 성경 근거 연결", style="Heading 2")
    doc.add_paragraph(f"명시 근거 연결 {citation.get('mapped_count',0)}건 · 근거 확인 필요 {citation.get('unsupported_count',0)}건")
    for item in citation.get("unsupported_claims", [])[:30]:
        doc.add_paragraph(f"[확인 필요 · 문장 {item.get('sentence')}] {item.get('text','')}")
    footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; footer.add_run("LM Studio 로컬 AI · " + ("목회자 승인 최종 잠금본" if review.get("state") == "locked" else "설교 초안 · 목회자 검토 필요")).font.size=Pt(8)
    doc.save(path)


def write_hwpx(path: Path, *, sermon: str, meta: dict) -> None:
    """Create a portable HWPX document from the vendored OWPML base template."""
    template = RESOURCE_ROOT / "app" / "assets" / "hwpx_base"
    section_path = template / "Contents" / "section0.xml"
    if not section_path.exists():
        raise RuntimeError("HWPX 기본 템플릿을 찾을 수 없습니다.")
    escaped = lambda value: html.escape(str(value or ""), quote=False)
    title = escaped(meta.get("topic") or "설교문")
    reference = escaped(meta.get("main_reference") or "")
    lines = [title, f"중심본문: {reference}" if reference else "", ""] + sermon_with_media_prompts(sermon, meta).splitlines()
    paragraphs = []
    for index, line in enumerate(lines, start=1000000100):
        text = line.strip()
        if text.startswith("### "):
            text = text[4:]
        elif text.startswith("## "):
            text = text[3:]
        elif text.startswith("# "):
            text = text[2:]
        content = f"<hp:t>{escaped(text)}</hp:t>" if text else "<hp:t/>"
        paragraphs.append(
            f'<hp:p id="{index}" paraPrIDRef="0" styleIDRef="0" pageBreak="0" columnBreak="0" merged="0">'
            f'<hp:run charPrIDRef="0">{content}</hp:run></hp:p>'
        )
    section = section_path.read_text(encoding="utf-8")
    section = section.replace("</hs:sec>", "\n" + "\n".join(paragraphs) + "\n</hs:sec>", 1)
    path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path, "w") as archive:
        # HWPX requires the mimetype entry to be the first, uncompressed entry.
        archive.writestr("mimetype", (template / "mimetype").read_bytes(), compress_type=zipfile.ZIP_STORED)
        for source in sorted(template.rglob("*")):
            if not source.is_file():
                continue
            name = source.relative_to(template).as_posix()
            if name == "mimetype":
                continue
            elif name == "Contents/section0.xml":
                archive.writestr(name, section.encode("utf-8"), compress_type=zipfile.ZIP_DEFLATED)
            else:
                archive.write(source, name, compress_type=zipfile.ZIP_DEFLATED)


def write_final_package(path: Path, *, sermon: str, meta: dict, sources: list[dict], project: dict) -> dict:
    package_meta = dict(meta)
    package_meta["project"] = dict(project)
    manifest = {
        "format": "sermon-lmstudio-final-package-v40",
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "topic": package_meta.get("topic", ""),
        "project": project,
        "target_minutes": package_meta.get("target_minutes"),
        "minutes_estimate": package_meta.get("minutes_estimate"),
        "reading_cpm": package_meta.get("reading_cpm", 330),
        "audit_status": (package_meta.get("audit") or {}).get("status"),
        "review_state": (package_meta.get("review_state") or {}).get("state"),
        "source_count": len(sources),
        "final_lock": (package_meta.get("review_state") or {}).get("lock"),
        "files": [],
    }
    path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as tmp:
        temp = Path(tmp)
        md_path = temp / "sermon.md"
        html_path = temp / "dashboard.html"
        docx_path = temp / "sermon.docx"
        pdf_path = temp / "sermon.pdf"
        md_path.write_text(sermon_with_media_prompts(sermon, package_meta), encoding="utf-8")
        media_path = temp / "media_prompts.md"
        media_path.write_text(media_prompts_markdown(package_meta.get("media_prompts")), encoding="utf-8")
        html_path.write_text(dashboard_html(sermon=sermon, meta=package_meta, sources=sources), encoding="utf-8")
        write_docx(docx_path, sermon=sermon, meta=package_meta)
        manifest["files"] = ["sermon.md", "dashboard.html", "sermon.docx", "media_prompts.md"]
        study = package_meta.get("study_note") if isinstance(package_meta.get("study_note"), dict) else {}
        if str(study.get("note_markdown", "")).strip():
            study_path = temp / "study_note.md"
            study_path.write_text(str(study["note_markdown"]), encoding="utf-8")
            manifest["files"].append("study_note.md")
        outline = package_meta.get("outline") if isinstance(package_meta.get("outline"), dict) else {}
        if outline:
            outline_path = temp / "sermon_outline.json"
            outline_path.write_text(json.dumps(outline, ensure_ascii=False, indent=2), encoding="utf-8")
            manifest["files"].append("sermon_outline.json")
        try:
            write_pdf(pdf_path, sermon=sermon, meta=package_meta, sources=sources)
            manifest["files"].append("sermon.pdf")
            manifest["pdf"] = {"included": True}
        except Exception as exc:
            # PDF is an optional companion artifact. A missing local font,
            # WeasyPrint/native library problem, or renderer-specific failure
            # must not prevent delivery of the locked source package.
            manifest["pdf"] = {"included": False, "reason": str(exc)}
        manifest["files"].append("manifest.json")
        manifest_path = temp / "manifest.json"
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for filename in manifest["files"]:
                archive.write(temp / filename, filename)
    return manifest
